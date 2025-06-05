import os
import sys
import re
from docx import Document
from docx.shared import Pt

def clean_text(text):
    # Word XMLに非対応な制御文字（NULLバイトなど）を削除
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)

def merge_documents(input_dir, base_filename, output_file, template_file=None):
    # テンプレートがあれば使用、なければ空の文書を作成
    if template_file and os.path.exists(template_file):
        merged_doc = Document(template_file)
    else:
        merged_doc = Document()

    part_num = 1
    while True:
        part_filename = os.path.join(input_dir, f"{base_filename}_part{part_num}_llm_output.txt")
        if not os.path.exists(part_filename):
            break

        with open(part_filename, "r", encoding="utf-8") as f:
            content = f.read()
            cleaned_content = clean_text(content)

        # セクションタイトル（Part番号）
        merged_doc.add_paragraph(f"Part {part_num}", "Heading 2")

        # 複数行に分けて段落を整える（空行で段落分け）
        for paragraph_text in cleaned_content.strip().split("\n\n"):
            paragraph = merged_doc.add_paragraph()
            for line in paragraph_text.strip().split("\n"):
                paragraph.add_run(clean_text(line)).font.size = Pt(11)

        part_num += 1

    merged_doc.save(output_file)
    print(f"✅ Merged document saved to {output_file}")

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python merge_docx.py <input_dir> <base_filename> <output_file> <template_file>")
        sys.exit(1)

    input_dir = sys.argv[1]
    base_filename = sys.argv[2]
    output_file = sys.argv[3]
    template_file = sys.argv[4]

    merge_documents(input_dir, base_filename, output_file, template_file)
