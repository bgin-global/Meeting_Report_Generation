import os
import re
import sys
from docx import Document

def clean_text(text):
    text = re.sub(r'\[\d+m', '', text)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    return text.strip()

def remove_executive_summary_section(doc):
    start_idx = None
    end_idx = None

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("1. Executive Summary"):
            start_idx = i
        elif para.text.strip().startswith("1.1 Condensed Executive Summary") and start_idx is not None:
            end_idx = i
            break

    if start_idx is not None and end_idx is not None:
        for i in range(end_idx - 1, start_idx, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

def replace_condensed_summary(doc, summary_text):
    for para in doc.paragraphs:
        if "<< EXEC_SUMMARY_CONDENSED >>" in para.text:
            for run in para.runs:
                if "<< EXEC_SUMMARY_CONDENSED >>" in run.text:
                    run.text = run.text.replace("<< EXEC_SUMMARY_CONDENSED >>", clean_text(summary_text))
            break

def refine_summary(input_docx_path, output_docx_path, new_summary_path):
    print(f"[refine_exec_summary] input_docx: {input_docx_path}")
    print(f"[refine_exec_summary] new_summary_path: {new_summary_path}")
    print(f"[refine_exec_summary] output_docx: {output_docx_path}")
    doc = Document(input_docx_path)
    with open(new_summary_path, "r", encoding="utf-8") as f:
        new_summary = f.read()
    print(f"[refine_exec_summary] new_summary: {new_summary}")
    remove_executive_summary_section(doc)
    replace_condensed_summary(doc, new_summary)
    doc.save(output_docx_path)
    print(f"✅ Refined report saved to: {output_docx_path}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python refine_exec_summary.py <input.docx> <new_summary.txt> <output.docx>")
        sys.exit(1)

    input_docx = sys.argv[1]
    summary_txt = sys.argv[2]
    output_docx = sys.argv[3]
    refine_summary(input_docx, output_docx, summary_txt)
