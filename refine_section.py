import sys
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def clean_text(text):
    text = re.sub(r'\[\d+m', '', text)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    return text.strip()

def find_section_indices_loose(doc, section_name):
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        if section_name in para.text.strip():
            print(f"[DEBUG] Found heading candidate at para {i}: {para.text.strip()}")
            start_idx = i
            continue
        if start_idx is not None:
            # 次の大見出し（数字+ピリオド or <<...>> プレースホルダー）で終わり
            if re.match(r"^\d+\. ", para.text.strip()) or re.match(r"^<<.*>>$", para.text.strip()):
                end_idx = i
                break
    if start_idx is not None and end_idx is None:
        end_idx = len(doc.paragraphs)
    print(f"[DEBUG] start_idx={start_idx}, end_idx={end_idx}")
    return start_idx, end_idx

def remove_section(doc, section_name):
    # Key Discussion Points, Action Items and Next Stepsは見出し番号付きで厳密検出
    section_headings = {
        "Key Discussion Points": ("2. Key Discussion Points", "2.1 Condensed Key Discussion Points"),
        "Action Items and Next Steps": ("3. Action Items and Next Steps", "3.1 Condensed Action Items and Next Steps")
    }
    if section_name in section_headings:
        start_heading, end_heading = section_headings[section_name]
        start_idx = None
        end_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith(start_heading):
                start_idx = i
            elif para.text.strip().startswith(end_heading) and start_idx is not None:
                end_idx = i
                break
        print(f"[DEBUG] start_idx={start_idx}, end_idx={end_idx}")
        if start_idx is not None and end_idx is not None:
            for i in range(end_idx - 1, start_idx, -1):
                print(f"[DEBUG] Removing: {doc.paragraphs[i].text}")
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)
        heading_para = doc.paragraphs[start_idx] if start_idx is not None else None
        return heading_para
    # それ以外は従来通り
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        if section_name in para.text.strip():
            start_idx = i
        elif start_idx is not None and (
            re.match(r"^\d+\.\d+ ", para.text.strip()) or  # 2.1 ... など
            re.match(r"^\d+\. ", para.text.strip()) or     # 3. ... など
            re.match(r"^<<.*>>$", para.text.strip())
        ):
            end_idx = i
            break
    print(f"[DEBUG] start_idx={start_idx}, end_idx={end_idx}")
    if start_idx is not None and end_idx is not None:
        for i in range(end_idx - 1, start_idx, -1):
            print(f"[DEBUG] Removing: {doc.paragraphs[i].text}")
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
    heading_para = doc.paragraphs[start_idx] if start_idx is not None else None
    return heading_para

def remove_original_placeholder(doc, section_name):
    placeholder_map = {
        "Key Discussion Points": "<< KEY_DISCUSSION >>",
        "Action Items and Next Steps": "<< ACTION_ITEMS >>"
    }
    placeholder = placeholder_map.get(section_name)
    if not placeholder:
        return
    for para in doc.paragraphs:
        if para.text.strip() == placeholder:
            # 段落ごと削除
            para._element.getparent().remove(para._element)
            return
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "")

def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para

def replace_condensed_placeholder(doc, section_name, summary_text, heading_para=None):
    placeholder_map = {
        "Key Discussion Points": "<< KEY_DISCUSSION_CONDENSED >>",
        "Action Items and Next Steps": "<< ACTION_ITEMS_CONDENSED >>"
    }
    placeholder = placeholder_map.get(section_name)
    found = False
    for para in doc.paragraphs:
        if placeholder and placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, clean_text(summary_text))
            found = True
            break
    # プレースホルダーが無い場合は、見出し直後に必ず挿入
    if not found and heading_para is not None:
        insert_paragraph_after(heading_para, clean_text(summary_text))

def refine_section(input_docx_path, output_docx_path, section_name, new_summary_path):
    doc = Document(input_docx_path)
    with open(new_summary_path, "r", encoding="utf-8") as f:
        new_summary = f.read()
    heading_para = remove_section(doc, section_name)
    remove_original_placeholder(doc, section_name)
    replace_condensed_placeholder(doc, section_name, new_summary, heading_para)
    doc.save(output_docx_path)
    print(f"✅ Refined report saved to: {output_docx_path}")

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python refine_section.py <input.docx> <section_name> <new_summary.txt> <output.docx>")
        sys.exit(1)
    input_docx = sys.argv[1]
    section_name = sys.argv[2]
    summary_txt = sys.argv[3]
    output_docx = sys.argv[4]

    # デバッグ: 段落・run構造を出力
    doc = Document(input_docx)
    print("[DEBUG] Paragraphs and runs in the input docx:")
    for i, para in enumerate(doc.paragraphs):
        print(f"[para {i}] {para.text}")
        for j, run in enumerate(para.runs):
            print(f"    [run {j}] {run.text}")
    # 通常処理
    refine_section(input_docx, output_docx, section_name, summary_txt) 